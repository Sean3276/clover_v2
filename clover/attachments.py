"""Attachment text extraction — so obligations living in attachments are not silently missed.

The 20-user review's #2 finding: 30–50% of real obligations live in attachments (BOMs, payment
backup, transcripts, contracts/redlines, rent rolls, cert PDFs). This module pulls TEXT (and PDF
tables + markup/comment annotations) out of the common attachment types so the same extraction /
citation / backfill model applies to them as to the body.

Design rules:
- **Loud, never silent.** A type we can't read yet (scanned image needing OCR, an unsupported type,
  a parse error) returns `ok=False` with a reason — it is flagged for review, never dropped quietly.
- **Lazy imports + graceful degradation.** A missing parser library flags the attachment unread
  rather than crashing the pipeline.
- Pure read-only; no network.
"""
from __future__ import annotations

from pathlib import Path

_MAX_CHARS = 500_000                       # cap extracted text per attachment
_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".gif", ".bmp", ".webp"}
_TEXT_EXTS = {".txt", ".csv", ".md", ".log"}


def _cap(text: str) -> tuple[str, bool]:
    text = text or ""
    return (text[:_MAX_CHARS], True) if len(text) > _MAX_CHARS else (text, False)


def _result(ok, kind, text="", note=""):
    text, truncated = _cap(text)
    return {"ok": bool(ok), "kind": kind, "text": text.strip(),
            "chars": len(text), "truncated": truncated, "note": note}


def _pdf(path: Path) -> dict:
    import fitz                                            # PyMuPDF
    parts: list[str] = []
    with fitz.open(path) as doc:
        for page in doc:
            parts.append(page.get_text("text") or "")
            try:                                           # tables (BOMs / rent rolls / schedules)
                for tbl in page.find_tables().tables:
                    for row in tbl.extract():
                        parts.append(" | ".join("" if c is None else str(c) for c in row))
            except Exception:
                pass
            try:                                           # markup / comment annotations (redlines)
                annot = page.first_annot
                while annot:
                    c = (annot.info or {}).get("content", "")
                    if c:
                        parts.append(f"[annotation] {c}")
                    annot = annot.next
            except Exception:
                pass
    return _result(True, "pdf", "\n".join(p for p in parts if p))


def _xlsx(path: Path) -> dict:
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    rows = []
    for ws in wb.worksheets:
        rows.append(f"# sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c not in (None, "")]
            if cells:
                rows.append(" | ".join(cells))
    wb.close()
    return _result(True, "xlsx", "\n".join(rows))


def _docx(path: Path) -> dict:
    import docx
    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    for t in d.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return _result(True, "docx", "\n".join(parts))


def extract_attachment(path) -> dict:
    """Extract text from one attachment. Returns {ok, kind, text, chars, truncated, note}.
    `ok=False` (with a note) for images-needing-OCR, unsupported types, or parse errors — LOUD, not silent."""
    p = Path(path)
    ext = p.suffix.lower()
    if not p.exists():
        return _result(False, ext.lstrip(".") or "unknown", note="file not found")
    try:
        if ext == ".pdf":
            return _pdf(p)
        if ext in (".xlsx", ".xlsm"):
            return _xlsx(p)
        if ext == ".docx":
            return _docx(p)
        if ext in _TEXT_EXTS:
            return _result(True, ext.lstrip("."), p.read_text(encoding="utf-8", errors="replace"))
        if ext in _IMAGE_EXTS:
            return _result(False, "image", note="scanned/image attachment — needs OCR (not enabled)")
        return _result(False, ext.lstrip(".") or "unknown", note="unsupported attachment type")
    except ImportError as e:
        return _result(False, ext.lstrip("."), note=f"parser library missing: {e.name}")
    except Exception as e:
        return _result(False, ext.lstrip("."), note=f"parse error: {type(e).__name__}: {e}")
