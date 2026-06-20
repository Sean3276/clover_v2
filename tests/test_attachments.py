from clover import attachments as att


def test_xlsx_extracts_cells(tmp_path):
    import openpyxl
    wb = openpyxl.Workbook(); ws = wb.active
    ws.append(["Ref", "Amount"]); ws.append(["INV-2024-001", 5000])
    p = tmp_path / "backup.xlsx"; wb.save(p)
    r = att.extract_attachment(p)
    assert r["ok"] and r["kind"] == "xlsx"
    assert "INV-2024-001" in r["text"] and "5000" in r["text"]


def test_docx_extracts_paragraphs_and_tables(tmp_path):
    import docx
    d = docx.Document(); d.add_paragraph("Action: submit RFI-12 by 14 Mar 2026")
    t = d.add_table(rows=1, cols=2); t.rows[0].cells[0].text = "PO"; t.rows[0].cells[1].text = "4500001"
    p = tmp_path / "sow.docx"; d.save(p)
    r = att.extract_attachment(p)
    assert r["ok"] and "RFI-12" in r["text"] and "4500001" in r["text"]


def test_pdf_extracts_text_and_annotation(tmp_path):
    import fitz
    doc = fitz.open(); page = doc.new_page()
    page.insert_text((72, 72), "PO 4500001 due 2026-06-30")
    page.add_text_annot((200, 200), "please review the redline by Friday")
    p = tmp_path / "contract.pdf"; doc.save(p); doc.close()
    r = att.extract_attachment(p)
    assert r["ok"] and "PO 4500001" in r["text"]
    assert "please review the redline" in r["text"]          # markup/comment annotation captured


def test_image_flagged_for_ocr_not_dropped(tmp_path):
    from PIL import Image
    p = tmp_path / "scan.png"; Image.new("RGB", (10, 10), "white").save(p)
    r = att.extract_attachment(p)
    assert r["ok"] is False and "OCR" in r["note"]           # loud, not silent


def test_unsupported_and_missing_are_loud(tmp_path):
    (tmp_path / "x.zip").write_bytes(b"PK\x03\x04")
    assert att.extract_attachment(tmp_path / "x.zip")["ok"] is False
    assert att.extract_attachment(tmp_path / "nope.pdf")["ok"] is False
